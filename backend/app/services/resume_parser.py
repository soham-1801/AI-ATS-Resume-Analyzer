import os
import pdfplumber
import logging
from docx import Document

logger = logging.getLogger("ats.resume_parser")

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from a PDF file using pdfplumber."""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for idx, page in enumerate(pdf.pages):
                    if idx >= 10:
                        break
                    page_text = page.extract_text()
                    if page_text:
                        if idx > 0:
                            text += "\f"
                        text += page_text + "\n"
        except Exception as e:
            err_msg = f"Corrupted document: PDF structure is invalid. Details: {str(e)}"
            logger.error(f"[RESUME PARSER] Error reading PDF {file_path}: {e}")
            raise ValueError(err_msg)
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from a Word document (.docx) using python-docx with docx2txt fallback.
        Supports paragraph, table, header, footer, and textbox extraction with diagnostics.
        """
        import zipfile
        import docx2txt

        filename = os.path.basename(file_path)
        ext = file_path.split(".")[-1].lower()
        file_size = os.path.getsize(file_path)
        
        print(f"[RESUME PARSER] Starting Word doc extraction. Filename: {filename}, Ext: {ext}, Size: {file_size} bytes")
        
        # Guard for .doc extension
        if ext == "doc":
            err_msg = "Unsupported file type: Legacy Word document (.doc) is not supported. Please save the document as a modern Word (.docx) file and re-upload."
            logger.error(f"[RESUME PARSER] Legacy Word document (.doc) upload attempted for: {filename}")
            raise ValueError(err_msg)
            
        # Basic validation: Is it a valid ZIP?
        is_zip = zipfile.is_zipfile(file_path)
        
        # Diagnostic variables
        pydocx_paragraphs = 0
        pydocx_tables = 0
        pydocx_textboxes = 0
        pydocx_sections = 0
        pydocx_char_count = 0
        
        docx2txt_char_count = 0
        
        primary_text = ""
        fallback_text = ""
        
        # 1. Primary Parser: python-docx
        try:
            print("[RESUME PARSER] Attempting primary extraction with python-docx...")
            doc = Document(file_path)
            pydocx_sections = len(doc.sections)
            
            extracted_chunks = []
            
            # A. Paragraphs
            for p in doc.paragraphs:
                pydocx_paragraphs += 1
                txt = p.text.strip()
                if txt:
                    extracted_chunks.append(txt)
                    
            # B. Tables
            for table in doc.tables:
                pydocx_tables += 1
                for row in table.rows:
                    row_text = []
                    seen_cells = set()
                    for cell in row.cells:
                        # Dedup cell XML element
                        if cell._tc in seen_cells:
                            continue
                        seen_cells.add(cell._tc)
                        cell_vals = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
                        if cell_vals:
                            row_text.append(" | ".join(cell_vals))
                    if row_text:
                        extracted_chunks.append(" ".join(row_text))
                        
            # C. Headers & Footers
            for section in doc.sections:
                # Headers
                headers = [section.header, section.first_page_header, section.even_page_header]
                for h in headers:
                    if h:
                        for p in h.paragraphs:
                            txt = p.text.strip()
                            if txt:
                                extracted_chunks.append(txt)
                        for t in h.tables:
                            for r in t.rows:
                                for cell in r.cells:
                                    txt = cell.text.strip()
                                    if txt:
                                        extracted_chunks.append(txt)
                # Footers
                footers = [section.footer, section.first_page_footer, section.even_page_footer]
                for f in footers:
                    if f:
                        for p in f.paragraphs:
                            txt = p.text.strip()
                            if txt:
                                extracted_chunks.append(txt)
                        for t in f.tables:
                            for r in t.rows:
                                for cell in r.cells:
                                    txt = cell.text.strip()
                                    if txt:
                                        extracted_chunks.append(txt)

            # D. Text boxes (drawing / w:txbxContent)
            try:
                txbx_elements = doc.element.xpath('//w:txbxContent')
                for txbx in txbx_elements:
                    from docx.text.paragraph import Paragraph
                    for p_elem in txbx.xpath('.//w:p'):
                        pydocx_textboxes += 1
                        p = Paragraph(p_elem, doc)
                        txt = p.text.strip()
                        if txt:
                            extracted_chunks.append(txt)
            except Exception as txbx_err:
                print(f"[RESUME PARSER] Warning: python-docx text box xpath extraction failed: {txbx_err}")
                
            # Combine primary text
            primary_text = "\n".join(extracted_chunks).strip()
            pydocx_char_count = len(primary_text)
            print(f"[RESUME PARSER] python-docx extracted {pydocx_char_count} characters.")
            
        except Exception as pydocx_err:
            print(f"[RESUME PARSER] python-docx parser failed. Error: {pydocx_err}")
            primary_text = ""

        # 2. Fallback check: Retry with docx2txt if primary returns empty text or < 20 characters
        if not primary_text or pydocx_char_count < 20:
            retry_reason = "empty text" if not primary_text else f"insufficient text length ({pydocx_char_count} chars < 20)"
            print(f"[RESUME PARSER] Primary parser returned {retry_reason}. Retrying with docx2txt fallback...")
            
            try:
                fallback_text = docx2txt.process(file_path)
                if fallback_text:
                    fallback_text = fallback_text.strip()
                    docx2txt_char_count = len(fallback_text)
                print(f"[RESUME PARSER] docx2txt fallback extracted {docx2txt_char_count} characters.")
            except Exception as docx2txt_err:
                print(f"[RESUME PARSER] docx2txt fallback also failed. Error: {docx2txt_err}")
                fallback_text = ""

        # Choose the best text
        final_text = primary_text if len(primary_text) >= len(fallback_text) else fallback_text
        final_char_count = len(final_text)
        
        # Display Diagnostics in logs
        diagnostics = (
            f"\n[PARSER DIAGNOSTICS] File: {filename}\n"
            f"- File Size: {file_size} bytes\n"
            f"- Is Valid ZIP: {is_zip}\n"
            f"- python-docx results:\n"
            f"  - Sections: {pydocx_sections}\n"
            f"  - Paragraphs: {pydocx_paragraphs}\n"
            f"  - Tables: {pydocx_tables}\n"
            f"  - Text Boxes parsed: {pydocx_textboxes}\n"
            f"  - Extracted Chars: {pydocx_char_count}\n"
            f"- docx2txt results:\n"
            f"  - Extracted Chars: {docx2txt_char_count}\n"
            f"- Final Chosen Chars: {final_char_count}\n"
        )
        print(diagnostics)
        logger.info(diagnostics)

        # 3. If extracted text length < 20 characters, return detailed diagnostics in exception
        if final_char_count < 20:
            if not is_zip:
                detailed_msg = (
                    f"Corrupted document: The file structure is invalid or corrupted (not a valid ZIP/DOCX structure).\n"
                    f"Diagnostics:\n"
                    f"- File size: {file_size} bytes\n"
                    f"- Is valid DOCX ZIP structure: {is_zip}\n"
                    f"Please verify the document works in MS Word."
                )
            else:
                detailed_msg = (
                    f"Empty document: The file contains no readable text content (extracted {final_char_count} characters).\n"
                    f"Diagnostics:\n"
                    f"- File size: {file_size} bytes\n"
                    f"- Is valid DOCX ZIP structure: {is_zip}\n"
                    f"- Paragraphs found: {pydocx_paragraphs}\n"
                    f"- Tables found: {pydocx_tables}\n"
                    f"- Text boxes found: {pydocx_textboxes}\n"
                    f"Please ensure the document contains readable text and is not an image-only scan or empty layout."
                )
            logger.error(f"[RESUME PARSER] File extraction failed: {detailed_msg}")
            raise ValueError(detailed_msg)

        # Log character count before ATS analysis (also logged globally at api layer)
        logger.info(f"[RESUME PARSER] Successfully extracted {final_char_count} characters from {filename}")
        return final_text

    @classmethod
    def run_ocr_on_pdf(cls, file_path: str) -> str:
        """Runs OCR on a scanned PDF using pytesseract with a high-fidelity Gemini API fallback."""
        print(f"[RESUME PARSER] Scanned PDF detected. Initiating OCR pipeline for: {file_path}")
        extracted_text = ""
        
        # 1. Attempt local OCR via pytesseract
        try:
            import pytesseract
            # Render PDF pages to images using pdfplumber (uses pdfium natively)
            with pdfplumber.open(file_path) as pdf:
                for idx, page in enumerate(pdf.pages):
                    if idx >= 5:
                        break
                    print(f"[RESUME PARSER] Rendering page {idx + 1} to image...")
                    # Convert page to high-res PIL image
                    pil_img = page.to_image(resolution=150).original
                    # Perform OCR
                    page_text = pytesseract.image_to_string(pil_img)
                    if page_text:
                        if idx > 0:
                            extracted_text += "\f"
                        extracted_text += page_text + "\n"
            
            if extracted_text.strip():
                print(f"[RESUME PARSER] Local pytesseract OCR succeeded. Extracted {len(extracted_text)} characters.")
                return extracted_text
                
        except Exception as local_err:
            print(f"[RESUME PARSER] Local pytesseract OCR failed or Tesseract is not installed. Error: {local_err}")
        
        # 2. Fallback to Gemini API-based multimodal PDF OCR
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            print("[RESUME PARSER] Attempting Gemini API multimodal PDF OCR fallback...")
            try:
                from google import genai
                from google.genai import types
                client = genai.Client(api_key=api_key)
                
                with open(file_path, "rb") as f:
                    pdf_bytes = f.read()
                
                part = types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")
                
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[part, "Perform high-fidelity OCR on this document. Extract all text contents exactly as they appear, preserving formatting and structure. Do not add any conversational intros or explanations."]
                )
                
                gemini_text = response.text
                if gemini_text and gemini_text.strip():
                    print(f"[RESUME PARSER] Gemini API OCR succeeded. Extracted {len(gemini_text)} characters.")
                    return gemini_text
            except Exception as gemini_err:
                print(f"[RESUME PARSER] Gemini API OCR fallback failed: {gemini_err}")
                
        # 3. If both fail, raise exception
        raise ValueError("Unable to extract text from the scanned PDF.")

    @classmethod
    def run_ocr_on_image(cls, file_path: str) -> str:
        """Runs OCR on an image file using pytesseract with a high-fidelity Gemini API fallback."""
        print(f"[RESUME PARSER] Image detected. Initiating OCR pipeline for: {file_path}")
        extracted_text = ""
        
        # 1. Attempt local OCR via pytesseract
        try:
            import pytesseract
            from PIL import Image
            print("[RESUME PARSER] Reading image...")
            pil_img = Image.open(file_path)
            # Perform OCR
            extracted_text = pytesseract.image_to_string(pil_img)
            
            if extracted_text and extracted_text.strip():
                print(f"[RESUME PARSER] Local pytesseract OCR succeeded on image. Extracted {len(extracted_text)} characters.")
                return extracted_text
                
        except Exception as local_err:
            print(f"[RESUME PARSER] Local pytesseract OCR failed on image or Tesseract is not installed. Error: {local_err}")
        
        # 2. Fallback to Gemini API-based multimodal OCR
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            print("[RESUME PARSER] Attempting Gemini API multimodal image OCR fallback...")
            try:
                from google import genai
                from google.genai import types
                client = genai.Client(api_key=api_key)
                
                with open(file_path, "rb") as f:
                    image_bytes = f.read()
                
                ext = file_path.rsplit(".", 1)[-1].lower()
                mime_type = "image/png" if ext == "png" else "image/jpeg"

                part = types.Part.from_bytes(data=image_bytes, mime_type=mime_type)
                
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=[part, "Perform high-fidelity OCR on this document. Extract all text contents exactly as they appear, preserving formatting and structure. Do not add any conversational intros or explanations."]
                )
                
                gemini_text = response.text
                if gemini_text and gemini_text.strip():
                    print(f"[RESUME PARSER] Gemini API OCR succeeded on image. Extracted {len(gemini_text)} characters.")
                    return gemini_text
            except Exception as gemini_err:
                print(f"[RESUME PARSER] Gemini API OCR fallback on image failed: {gemini_err}")
                
        # 3. If both fail, raise exception
        raise ValueError("Unable to extract text from the image file.")

    @classmethod
    def parse(cls, file_path: str) -> str:
        """Main parsing router based on file extension."""
        print(f"[RESUME PARSER] Starting parsing for file: {file_path}")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at: {file_path}")
            
        ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""
        if ext == "pdf":
            text = cls.extract_text_from_pdf(file_path)
            if not text.strip():
                print(f"[RESUME PARSER] WARNING: No readable text extracted via pdfplumber. Redirecting to OCR pipeline for {file_path}...")
                text = cls.run_ocr_on_pdf(file_path)
        elif ext in ["docx", "doc"]:
            text = cls.extract_text_from_docx(file_path)
        elif ext in ["jpg", "jpeg", "png"]:
            text = cls.run_ocr_on_image(file_path)
        elif ext == "txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                if not text.strip():
                    raise ValueError("The text file is empty.")
            except ValueError:
                raise
            except Exception as e:
                print(f"[RESUME PARSER] Error reading TXT {file_path}: {e}")
                raise ValueError(f"Failed to read TXT file: {e}")
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
        text_len = len(text) if text else 0
        print(f"[RESUME PARSER] Extraction complete. File: {os.path.basename(file_path)}, Text Length: {text_len} characters")
        return text
