import requests
import time
import subprocess
import os
import uuid
from fpdf import FPDF
from docx import Document

PORT = 8004
BASE_URL = f"http://127.0.0.1:{PORT}/api"

def create_files():
    # 1. Empty PDF (creates a PDF with 0 extractable text)
    pdf = FPDF()
    pdf.add_page()
    # Add no text
    pdf.output("empty_test.pdf")

    # 2. Invalid File (Corrupt/Fake PDF)
    with open("invalid_test.pdf", "w") as f:
        f.write("This is not a real PDF file, just random text.")

    # 3. Valid PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="John Doe - Software Engineer", ln=True, align='C')
    pdf.cell(200, 10, txt="Email: john@example.com", ln=True, align='C')
    pdf.cell(200, 10, txt="Experience: 5 years in Python and React.", ln=True, align='L')
    pdf.output("valid_test.pdf")

    # 4. DOCX
    doc = Document()
    doc.add_heading('Jane Smith - Software Engineer', 0)
    doc.add_paragraph('Email: jane@example.com')
    doc.add_paragraph('Experience: 5 years in Python and React.')
    doc.save("valid_test.docx")

    # 5. TXT
    with open("valid_test.txt", "w") as f:
        f.write("Alice Brown - Data Scientist\nEmail: alice@example.com\nSkills: Python, Machine Learning, SQL.")

def cleanup():
    for f in ["empty_test.pdf", "invalid_test.pdf", "valid_test.pdf", "valid_test.docx", "valid_test.txt"]:
        if os.path.exists(f):
            os.remove(f)

def run_tests():
    create_files()
    
    print("--- Starting Upload Tests ---")
    env = os.environ.copy()
    python_exe = os.path.join(os.path.abspath(os.path.dirname(__file__)), "venv", "Scripts", "python.exe")
    server_process = subprocess.Popen(
        [python_exe, "-m", "uvicorn", "app.main:app", "--port", str(PORT)],
        cwd=os.path.abspath(os.path.dirname(__file__)),
        env=env
    )
    
    # Wait for server
    server_up = False
    for _ in range(20):
        try:
            r = requests.get(f"http://127.0.0.1:{PORT}")
            if r.status_code == 200:
                server_up = True
                break
        except requests.exceptions.ConnectionError:
            time.sleep(0.5)
            
    if not server_up:
        print("Failed to start server")
        server_process.kill()
        return

    try:
        # Register and login
        uid = str(uuid.uuid4())[:8]
        email = f"uploader_{uid}@example.com"
        password = "Password123!"
        
        requests.post(f"{BASE_URL}/auth/register", json={"email": email, "password": password, "name": "Upload Test User"})
        r_login = requests.post(f"{BASE_URL}/auth/login", json={"email": email, "password": password})
        token = r_login.json().get("access_token")
        headers = {"Authorization": f"Bearer {token}"}
        
        def upload_file(filepath, mime_type):
            with open(filepath, "rb") as f:
                files = {"file": (os.path.basename(filepath), f, mime_type)}
                r = requests.post(f"{BASE_URL}/resume/upload", headers=headers, files=files)
                return r

        print("\n1. Testing Empty PDF")
        r1 = upload_file("empty_test.pdf", "application/pdf")
        print(f"Status: {r1.status_code} - {r1.text}")
        
        print("\n2. Testing Invalid File (Fake PDF)")
        r2 = upload_file("invalid_test.pdf", "application/pdf")
        print(f"Status: {r2.status_code} - {r2.text}")
        
        print("\n3. Testing Valid PDF")
        r3 = upload_file("valid_test.pdf", "application/pdf")
        print(f"Status: {r3.status_code} - " + ("Success" if r3.status_code == 201 else r3.text))

        print("\n4. Testing Valid DOCX")
        r4 = upload_file("valid_test.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        print(f"Status: {r4.status_code} - " + ("Success" if r4.status_code == 201 else r4.text))

        print("\n5. Testing Valid TXT")
        r5 = upload_file("valid_test.txt", "text/plain")
        print(f"Status: {r5.status_code} - " + ("Success" if r5.status_code == 201 else r5.text))

    except Exception as e:
        print(f"Error: {e}")
    finally:
        server_process.kill()
        cleanup()
        print("\n--- Upload Tests Finished ---")

if __name__ == "__main__":
    run_tests()
